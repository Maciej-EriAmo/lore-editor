"""Eksport rękopisu do DOCX (format wydawniczy)."""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from lore.manuscript import ManuscriptProfile, get_profile


class DocxExportError(RuntimeError):
    pass


def _require_docx():
    try:
        from docx import Document
        from docx.shared import Cm, Pt
    except ImportError as e:
        raise DocxExportError(
            "Brak pakietu python-docx. Zainstaluj: pip install python-docx"
        ) from e
    return Document, Cm, Pt


def export_manuscript_docx(
    text: str,
    path: str | Path,
    *,
    profile_id: str = "submission",
    title: str = "",
) -> Path:
    """Zapisz tekst jako .docx z marginesami i czcionką profilu wydawniczego."""
    Document, Cm, Pt = _require_docx()
    profile = get_profile(profile_id)
    dest = Path(path)
    dest.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    margin = Cm(profile.margin_mm / 10.0)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    normal = doc.styles["Normal"]
    normal.font.name = profile.font_family
    normal.font.size = Pt(profile.font_size_pt)
    pf = normal.paragraph_format
    pf.line_spacing = profile.line_spacing

    if title:
        h = doc.add_heading(title, level=1)
        h.runs[0].font.name = profile.font_family

    if profile.screenplay_timing or profile.lines_per_page:
        for line in text.splitlines():
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = profile.line_spacing
    else:
        blocks = text.split("\n\n")
        if len(blocks) <= 1:
            for line in text.splitlines():
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = profile.line_spacing
        else:
            for block in blocks:
                p = doc.add_paragraph(block.replace("\n", " "))
                p.paragraph_format.line_spacing = profile.line_spacing

    doc.save(str(dest))
    return dest


def export_available() -> bool:
    try:
        _require_docx()
        return True
    except DocxExportError:
        return False